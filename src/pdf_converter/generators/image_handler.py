"""Image sizing, embedding, and fallbacks for Word generation.

Handles loading images from file paths or base64, auto-scaling to
configured max dimensions, and placeholder generation when images
are unavailable.
"""

from __future__ import annotations

import base64
import io
import logging
from pathlib import Path
from typing import Optional

from docx.document import Document
from docx.shared import Inches

from pdf_converter.config import ImageConfig
from pdf_converter.exceptions import ImageError
from pdf_converter.ir.schema import FigureBlock

logger = logging.getLogger(__name__)


def add_image(
    doc: Document,
    block: FigureBlock,
    config: ImageConfig,
    base_dir: Optional[Path] = None,
) -> None:
    """Add an image from a FigureBlock to the document.

    Tries image_path first (relative to base_dir), then image_base64.
    Falls back to a placeholder paragraph if neither works.

    Args:
        doc: The python-docx Document.
        block: The IR FigureBlock with image data.
        config: Image configuration (max dimensions, placeholder text).
        base_dir: Base directory for resolving relative image paths.
    """
    image_stream = _load_image(block, base_dir)

    if image_stream is None:
        _add_placeholder(doc, config)
        return

    width, height = _compute_dimensions(image_stream, block, config)
    image_stream.seek(0)

    doc.add_picture(image_stream, width=width, height=height)


def _load_image(
    block: FigureBlock, base_dir: Optional[Path]
) -> Optional[io.BytesIO]:
    """Load image bytes from path or base64."""
    # Try file path first
    if block.image_path:
        path = Path(block.image_path)
        if base_dir and not path.is_absolute():
            path = base_dir / path

        if path.exists():
            try:
                return io.BytesIO(path.read_bytes())
            except OSError as exc:
                logger.warning("Failed to read image %s: %s", path, exc)

    # Try base64 fallback
    if block.image_base64:
        try:
            data = base64.b64decode(block.image_base64)
            return io.BytesIO(data)
        except Exception as exc:
            logger.warning("Failed to decode base64 image: %s", exc)

    return None


def _compute_dimensions(
    image_stream: io.BytesIO,
    block: FigureBlock,
    config: ImageConfig,
) -> tuple[Optional[Inches], Optional[Inches]]:
    """Compute scaled dimensions that fit within max bounds.

    Uses block-specified dimensions if available, otherwise reads
    the image to get its native size, then scales down to fit
    within config.max_width_inches × config.max_height_inches.
    """
    width_in = block.width_inches
    height_in = block.height_inches

    if width_in is None or height_in is None:
        try:
            from PIL import Image

            image_stream.seek(0)
            img = Image.open(image_stream)
            dpi_x, dpi_y = img.info.get("dpi", (96, 96))
            # Ensure DPI values are reasonable
            dpi_x = max(dpi_x, 72)
            dpi_y = max(dpi_y, 72)
            width_in = img.width / dpi_x
            height_in = img.height / dpi_y
        except Exception:
            # Can't determine size — use max width, no height constraint
            return Inches(config.max_width_inches), None

    # Scale down to fit within bounds
    scale = 1.0
    if width_in > config.max_width_inches:
        scale = min(scale, config.max_width_inches / width_in)
    if height_in > config.max_height_inches:
        scale = min(scale, config.max_height_inches / height_in)

    return Inches(width_in * scale), Inches(height_in * scale)


def _add_placeholder(doc: Document, config: ImageConfig) -> None:
    """Add a placeholder paragraph when an image is unavailable."""
    p = doc.add_paragraph()
    run = p.add_run(config.placeholder_text)
    run.italic = True
