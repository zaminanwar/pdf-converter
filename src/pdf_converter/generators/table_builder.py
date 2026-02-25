"""Table rendering with merged cell support for python-docx.

Converts the flat TableCell list from the IR into a python-docx Table
with properly merged cells (row_span/col_span).
"""

from __future__ import annotations

from docx.document import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table

from pdf_converter.config import Config
from pdf_converter.generators.styles import apply_highlight
from pdf_converter.ir.schema import TableBlock, TextRun


def build_table(doc: Document, block: TableBlock, config: Config) -> Table:
    """Create a python-docx Table from a TableBlock.

    Args:
        doc: The python-docx Document to add the table to.
        block: The IR TableBlock containing cell data.
        config: Application configuration.

    Returns:
        The created Table object.
    """
    num_rows = block.num_rows or _infer_rows(block)
    num_cols = block.num_cols or _infer_cols(block)

    if num_rows == 0 or num_cols == 0:
        # Empty table — add a minimal 1x1 placeholder
        return doc.add_table(rows=1, cols=1, style=config.style.table_style)

    table = doc.add_table(rows=num_rows, cols=num_cols, style=config.style.table_style)

    # Build a lookup from (row, col) → cell data
    cell_map = {(c.row, c.col): c for c in block.cells}

    # Track which cells have been merged (to avoid overwriting)
    merged = set()

    for cell_data in block.cells:
        r, c = cell_data.row, cell_data.col
        if (r, c) in merged:
            continue
        if r >= num_rows or c >= num_cols:
            continue

        # Merge cells if spanning
        end_row = min(r + cell_data.row_span - 1, num_rows - 1)
        end_col = min(c + cell_data.col_span - 1, num_cols - 1)

        cell = table.cell(r, c)

        if end_row > r or end_col > c:
            cell = cell.merge(table.cell(end_row, end_col))
            # Mark all spanned positions as merged
            for mr in range(r, end_row + 1):
                for mc in range(c, end_col + 1):
                    if (mr, mc) != (r, c):
                        merged.add((mr, mc))

        # Set cell text
        if cell_data.runs:
            _write_runs_to_cell(cell, cell_data.runs)
        else:
            cell.text = cell_data.text

    # Apply table-wide formatting
    _set_table_autofit(table)
    _set_header_row(table)

    return table


def _write_runs_to_cell(cell, runs: list[TextRun]) -> None:
    """Write formatted text runs into a table cell."""
    # Clear default paragraph
    paragraph = cell.paragraphs[0]
    paragraph.clear()

    for run_data in runs:
        run = paragraph.add_run(run_data.text)
        _apply_run_formatting(run, run_data)


def _apply_run_formatting(run, run_data: TextRun) -> None:
    """Apply all formatting attributes to a run."""
    run.bold = run_data.bold
    run.italic = run_data.italic
    run.underline = run_data.underline
    if run_data.strikethrough:
        run.font.strike = True
    if run_data.superscript:
        run.font.superscript = True
    if run_data.subscript:
        run.font.subscript = True
    if run_data.highlight:
        apply_highlight(run, run_data.highlight)


def _set_table_autofit(table: Table) -> None:
    """Set table width to 100% of page width via OOXML."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")  # 5000 = 100% in fifths of a percent
    tblPr.append(tblW)


def _set_header_row(table: Table) -> None:
    """Mark the first row as a header row so it repeats across pages."""
    if len(table.rows) == 0:
        return
    first_row = table.rows[0]
    trPr = first_row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


def _infer_rows(block: TableBlock) -> int:
    """Infer row count from cell data."""
    if not block.cells:
        return 0
    return max(c.row + c.row_span for c in block.cells)


def _infer_cols(block: TableBlock) -> int:
    """Infer column count from cell data."""
    if not block.cells:
        return 0
    return max(c.col + c.col_span for c in block.cells)
