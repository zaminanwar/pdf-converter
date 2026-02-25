"""Recursive IR tree → .docx renderer.

Walks the IR tree depth-first, generating Word content for each block type.
Headings establish the document hierarchy; their children are rendered
recursively beneath them.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt

from pdf_converter.config import Config
from pdf_converter.exceptions import GenerationError
from pdf_converter.generators.image_handler import add_image
from pdf_converter.generators.styles import (
    apply_caption_formatting,
    apply_highlight,
    apply_list_numbering,
    ensure_styles_exist,
    heading_style_name,
    list_style_name,
)
from pdf_converter.generators.table_builder import build_table
from pdf_converter.ir.schema import (
    DocumentIR,
    FigureBlock,
    HeadingBlock,
    IRBlock,
    ListBlock,
    ListItem,
    PageBreakBlock,
    ParagraphBlock,
    TableBlock,
    TextRun,
)

logger = logging.getLogger(__name__)


class WordGenerator:
    """Generates a Word document from a DocumentIR tree."""

    def __init__(self, config: Optional[Config] = None):
        self.config = config or Config.default()

    def generate(
        self,
        ir: DocumentIR,
        output_path: Path,
        base_dir: Optional[Path] = None,
    ) -> Path:
        """Generate a .docx file from a DocumentIR.

        Args:
            ir: The document IR to render.
            output_path: Where to write the .docx file.
            base_dir: Base directory for resolving relative image paths.
                      Defaults to output_path's parent directory.

        Returns:
            The output path (for convenience).
        """
        if base_dir is None:
            base_dir = output_path.parent

        doc = Document()
        ensure_styles_exist(doc, self.config.style)

        # Set document title from metadata
        if ir.metadata.title:
            doc.core_properties.title = ir.metadata.title

        for block in ir.body:
            self._render_block(doc, block, base_dir)

        try:
            doc.save(str(output_path))
        except OSError as exc:
            raise GenerationError(f"Failed to save document: {exc}") from exc

        logger.info("Generated %s", output_path)
        return output_path

    def generate_document(
        self,
        ir: DocumentIR,
        base_dir: Optional[Path] = None,
    ) -> Document:
        """Generate and return a python-docx Document object (for testing).

        Args:
            ir: The document IR to render.
            base_dir: Base directory for resolving relative image paths.

        Returns:
            The python-docx Document object.
        """
        doc = Document()
        ensure_styles_exist(doc, self.config.style)

        # Set document title from metadata
        if ir.metadata.title:
            doc.core_properties.title = ir.metadata.title

        for block in ir.body:
            self._render_block(doc, block, base_dir)

        return doc

    def _render_block(
        self, doc: Document, block: IRBlock, base_dir: Optional[Path]
    ) -> None:
        """Dispatch rendering to the appropriate method by block type."""
        if isinstance(block, HeadingBlock):
            self._render_heading(doc, block, base_dir)
        elif isinstance(block, ParagraphBlock):
            self._render_paragraph(doc, block)
        elif isinstance(block, ListBlock):
            self._render_list(doc, block)
        elif isinstance(block, TableBlock):
            self._render_table(doc, block)
        elif isinstance(block, FigureBlock):
            self._render_figure(doc, block, base_dir)
        elif isinstance(block, PageBreakBlock):
            self._render_page_break(doc)
        else:
            logger.warning("Unknown block type: %s", type(block).__name__)

    def _render_heading(
        self, doc: Document, block: HeadingBlock, base_dir: Optional[Path]
    ) -> None:
        """Render a heading and recursively render its children.

        Always uses doc.add_paragraph(style=...) to respect heading_prefix
        configuration, never doc.add_heading() which ignores it.
        """
        style_name = heading_style_name(self.config.style, block.level)
        paragraph = doc.add_paragraph(style=style_name)

        if block.runs:
            _write_runs(paragraph, block.runs)
        else:
            paragraph.add_run(block.text)

        # Mark low-confidence headings if enabled
        if (
            self.config.style.mark_low_confidence
            and block.confidence < self.config.style.low_confidence_threshold
        ):
            for run in paragraph.runs:
                apply_highlight(run, self.config.style.low_confidence_highlight)
            marker = paragraph.add_run(
                f"  [{block.confidence:.0%}]"
            )
            marker.font.size = Pt(8)
            marker.font.italic = True

        # Recursively render children
        for child in block.children:
            self._render_block(doc, child, base_dir)

    def _render_paragraph(self, doc: Document, block: ParagraphBlock) -> None:
        """Render a paragraph with optional formatted runs."""
        if block.runs:
            paragraph = doc.add_paragraph(style=self.config.style.body_style)
            _write_runs(paragraph, block.runs)
        else:
            doc.add_paragraph(block.text, style=self.config.style.body_style)

    def _render_list(self, doc: Document, block: ListBlock) -> None:
        """Render a list block with nested items."""
        ordered = block.style == "ordered"
        marker_format = block.marker_format
        for item in block.items:
            self._render_list_item(doc, item, ordered, level=1, marker_format=marker_format)

    def _render_list_item(
        self, doc: Document, item: ListItem, ordered: bool, level: int,
        marker_format: str | None = None,
    ) -> None:
        """Render a single list item, then recurse into children."""
        style = list_style_name(self.config.style, ordered, level)

        if item.runs:
            paragraph = doc.add_paragraph(style=style)
            _write_runs(paragraph, item.runs)
        else:
            paragraph = doc.add_paragraph(item.text, style=style)

        # Apply numbering XML so bullets/numbers actually render
        apply_list_numbering(paragraph, doc, ordered, level, marker_format=marker_format)

        for child in item.children:
            self._render_list_item(doc, child, ordered, level=min(level + 1, 3), marker_format=marker_format)

    def _render_table(self, doc: Document, block: TableBlock) -> None:
        """Render a table using the table builder."""
        build_table(doc, block, self.config)

    def _render_figure(
        self, doc: Document, block: FigureBlock, base_dir: Optional[Path]
    ) -> None:
        """Render a figure (image + optional caption)."""
        add_image(doc, block, self.config.image, base_dir)

        if block.caption:
            p = doc.add_paragraph(block.caption)
            apply_caption_formatting(p, self.config.style)

    def _render_page_break(self, doc: Document) -> None:
        """Render a page break."""
        from docx.enum.text import WD_BREAK

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)


def _write_runs(paragraph, runs: list[TextRun]) -> None:
    """Write formatted TextRuns into a paragraph."""
    for run_data in runs:
        run = paragraph.add_run(run_data.text)
        run.bold = run_data.bold
        run.italic = run_data.italic
        run.underline = run_data.underline
        if run_data.strikethrough:
            run.font.strike = True
        if run_data.superscript:
            run.font.superscript = True
        if run_data.subscript:
            run.font.subscript = True
        if run_data.highlight:
            apply_highlight(run, run_data.highlight)
