"""Tests for Word generator — hand-craft IR objects → generate .docx → inspect."""

from pathlib import Path

import pytest
from docx import Document as open_docx
from docx.document import Document as DocxDocument

from pdf_converter.config import Config
from pdf_converter.generators.word_generator import WordGenerator
from pdf_converter.ir.schema import (
    DocumentIR,
    DocumentMetadata,
    FigureBlock,
    HeadingBlock,
    ListBlock,
    ListItem,
    PageBreakBlock,
    ParagraphBlock,
    TableBlock,
    TableCell,
    TextRun,
)


@pytest.fixture
def generator():
    return WordGenerator()


@pytest.fixture
def simple_ir():
    """A simple IR with a heading and paragraph."""
    return DocumentIR(
        metadata=DocumentMetadata(title="Test Doc"),
        body=[
            HeadingBlock(level=1, text="Introduction", children=[
                ParagraphBlock(text="Hello world."),
            ]),
        ],
    )


# ---------------------------------------------------------------------------
# Basic generation
# ---------------------------------------------------------------------------

class TestBasicGeneration:
    def test_generates_docx_file(self, generator, simple_ir, tmp_path):
        out = tmp_path / "out.docx"
        result = generator.generate(simple_ir, out)
        assert result == out
        assert out.exists()
        assert out.stat().st_size > 0

    def test_generates_document_object(self, generator, simple_ir):
        doc = generator.generate_document(simple_ir)
        assert isinstance(doc, DocxDocument)

    def test_empty_document(self, generator):
        ir = DocumentIR()
        doc = generator.generate_document(ir)
        # Should produce a valid document with no body content
        assert len(doc.paragraphs) == 0


# ---------------------------------------------------------------------------
# Headings
# ---------------------------------------------------------------------------

class TestHeadings:
    def test_heading_text(self, generator):
        ir = DocumentIR(body=[
            HeadingBlock(level=1, text="Chapter 1"),
        ])
        doc = generator.generate_document(ir)
        assert doc.paragraphs[0].text == "Chapter 1"

    def test_heading_hierarchy(self, generator):
        ir = DocumentIR(body=[
            HeadingBlock(level=1, text="H1", children=[
                ParagraphBlock(text="Under H1"),
                HeadingBlock(level=2, text="H2", children=[
                    ParagraphBlock(text="Under H2"),
                ]),
            ]),
        ])
        doc = generator.generate_document(ir)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["H1", "Under H1", "H2", "Under H2"]

    def test_heading_levels(self, generator):
        ir = DocumentIR(body=[
            HeadingBlock(level=1, text="L1"),
            HeadingBlock(level=2, text="L2"),
            HeadingBlock(level=3, text="L3"),
        ])
        doc = generator.generate_document(ir)
        # Check style names include the level
        assert "Heading 1" in doc.paragraphs[0].style.name
        assert "Heading 2" in doc.paragraphs[1].style.name
        assert "Heading 3" in doc.paragraphs[2].style.name

    def test_heading_always_uses_configured_style(self, generator):
        """Headings must always use doc.add_paragraph(style=...), never doc.add_heading()."""
        ir = DocumentIR(body=[
            HeadingBlock(level=1, text="Plain heading"),
            HeadingBlock(level=2, text="With runs", runs=[
                TextRun(text="formatted", bold=True),
            ]),
        ])
        doc = generator.generate_document(ir)
        # Both should use Heading style
        assert "Heading 1" in doc.paragraphs[0].style.name
        assert "Heading 2" in doc.paragraphs[1].style.name


# ---------------------------------------------------------------------------
# Paragraphs
# ---------------------------------------------------------------------------

class TestParagraphs:
    def test_plain_text(self, generator):
        ir = DocumentIR(body=[ParagraphBlock(text="Plain paragraph.")])
        doc = generator.generate_document(ir)
        assert doc.paragraphs[0].text == "Plain paragraph."

    def test_formatted_runs(self, generator):
        ir = DocumentIR(body=[
            ParagraphBlock(
                text="",
                runs=[
                    TextRun(text="Normal "),
                    TextRun(text="bold", bold=True),
                    TextRun(text=" and "),
                    TextRun(text="italic", italic=True),
                ],
            ),
        ])
        doc = generator.generate_document(ir)
        para = doc.paragraphs[0]
        assert para.text == "Normal bold and italic"

        runs = para.runs
        assert runs[0].bold is not True
        assert runs[1].bold is True
        assert runs[3].italic is True

    def test_underline_and_strikethrough(self, generator):
        ir = DocumentIR(body=[
            ParagraphBlock(runs=[
                TextRun(text="underlined", underline=True),
                TextRun(text="struck", strikethrough=True),
            ]),
        ])
        doc = generator.generate_document(ir)
        runs = doc.paragraphs[0].runs
        assert runs[0].underline is True
        assert runs[1].font.strike is True


# ---------------------------------------------------------------------------
# Highlight
# ---------------------------------------------------------------------------

class TestHighlight:
    def test_highlight_run(self, generator):
        """A run with highlight=yellow should have w:highlight in the XML."""
        ir = DocumentIR(body=[
            ParagraphBlock(runs=[
                TextRun(text="highlighted", highlight="yellow"),
            ]),
        ])
        doc = generator.generate_document(ir)
        run = doc.paragraphs[0].runs[0]
        # Check that the highlight element is present in the run XML
        rPr = run._element.rPr
        assert rPr is not None
        highlights = rPr.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight"
        )
        assert len(highlights) == 1
        assert highlights[0].get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
        ) == "yellow"

    def test_no_highlight_when_none(self, generator):
        """A run without highlight should have no w:highlight element."""
        ir = DocumentIR(body=[
            ParagraphBlock(runs=[
                TextRun(text="normal"),
            ]),
        ])
        doc = generator.generate_document(ir)
        run = doc.paragraphs[0].runs[0]
        rPr = run._element.rPr
        if rPr is not None:
            highlights = rPr.findall(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}highlight"
            )
            assert len(highlights) == 0


# ---------------------------------------------------------------------------
# Document properties
# ---------------------------------------------------------------------------

class TestDocProperties:
    def test_title_set_from_metadata(self, generator):
        """doc.core_properties.title should be set from ir.metadata.title."""
        ir = DocumentIR(
            metadata=DocumentMetadata(title="My Document Title"),
            body=[ParagraphBlock(text="Content")],
        )
        doc = generator.generate_document(ir)
        assert doc.core_properties.title == "My Document Title"

    def test_no_title_when_empty(self, generator):
        """When metadata.title is empty, core_properties.title should not be set."""
        ir = DocumentIR(
            metadata=DocumentMetadata(title=""),
            body=[ParagraphBlock(text="Content")],
        )
        doc = generator.generate_document(ir)
        # Title should be empty/None (not set)
        assert not doc.core_properties.title

    def test_title_in_saved_file(self, generator, tmp_path):
        """Title should persist after saving and reopening."""
        ir = DocumentIR(
            metadata=DocumentMetadata(title="Saved Title"),
            body=[ParagraphBlock(text="Content")],
        )
        out = tmp_path / "titled.docx"
        generator.generate(ir, out)

        doc = open_docx(str(out))
        assert doc.core_properties.title == "Saved Title"


# ---------------------------------------------------------------------------
# Lists
# ---------------------------------------------------------------------------

class TestLists:
    def test_unordered_list(self, generator):
        ir = DocumentIR(body=[
            ListBlock(style="unordered", items=[
                ListItem(text="Item A"),
                ListItem(text="Item B"),
                ListItem(text="Item C"),
            ]),
        ])
        doc = generator.generate_document(ir)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["Item A", "Item B", "Item C"]
        # All should use List Bullet style
        for p in doc.paragraphs:
            assert "List Bullet" in p.style.name

    def test_ordered_list(self, generator):
        ir = DocumentIR(body=[
            ListBlock(style="ordered", items=[
                ListItem(text="First"),
                ListItem(text="Second"),
            ]),
        ])
        doc = generator.generate_document(ir)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["First", "Second"]
        for p in doc.paragraphs:
            assert "List Number" in p.style.name

    def test_nested_list(self, generator):
        ir = DocumentIR(body=[
            ListBlock(style="unordered", items=[
                ListItem(text="Parent", children=[
                    ListItem(text="Child 1"),
                    ListItem(text="Child 2"),
                ]),
            ]),
        ])
        doc = generator.generate_document(ir)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["Parent", "Child 1", "Child 2"]

    def test_list_has_numbering_xml(self, generator):
        """List items should have w:numPr XML for real bullet/number rendering."""
        ir = DocumentIR(body=[
            ListBlock(style="unordered", items=[
                ListItem(text="Bullet item"),
            ]),
        ])
        doc = generator.generate_document(ir)
        para = doc.paragraphs[0]
        pPr = para._element.pPr
        assert pPr is not None
        numPr = pPr.findall(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
        )
        assert len(numPr) >= 1

    def test_lower_letter_list_numbering(self, generator):
        """A list with marker_format='lowerLetter' should use numId 102."""
        ir = DocumentIR(body=[
            ListBlock(
                style="ordered",
                marker_format="lowerLetter",
                items=[
                    ListItem(text="item a"),
                    ListItem(text="item b"),
                ],
            ),
        ])
        doc = generator.generate_document(ir)
        para = doc.paragraphs[0]
        pPr = para._element.pPr
        assert pPr is not None
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        numPr = pPr.findall(f"{ns}numPr")
        assert len(numPr) >= 1
        numId_elem = numPr[0].findall(f"{ns}numId")
        assert len(numId_elem) >= 1
        assert numId_elem[0].get(f"{ns}val") == "102"

    def test_upper_letter_list_numbering(self, generator):
        """A list with marker_format='upperLetter' should use numId 103."""
        ir = DocumentIR(body=[
            ListBlock(
                style="ordered",
                marker_format="upperLetter",
                items=[ListItem(text="item A")],
            ),
        ])
        doc = generator.generate_document(ir)
        para = doc.paragraphs[0]
        pPr = para._element.pPr
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        numPr = pPr.findall(f"{ns}numPr")
        numId_elem = numPr[0].findall(f"{ns}numId")
        assert numId_elem[0].get(f"{ns}val") == "103"

    def test_default_ordered_list_numbering(self, generator):
        """A regular ordered list (no marker_format) should use numId 101."""
        ir = DocumentIR(body=[
            ListBlock(
                style="ordered",
                items=[ListItem(text="item 1")],
            ),
        ])
        doc = generator.generate_document(ir)
        para = doc.paragraphs[0]
        pPr = para._element.pPr
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        numPr = pPr.findall(f"{ns}numPr")
        numId_elem = numPr[0].findall(f"{ns}numId")
        assert numId_elem[0].get(f"{ns}val") == "101"


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------

class TestTables:
    def test_simple_table(self, generator):
        ir = DocumentIR(body=[
            TableBlock(
                num_rows=2,
                num_cols=2,
                cells=[
                    TableCell(row=0, col=0, text="A"),
                    TableCell(row=0, col=1, text="B"),
                    TableCell(row=1, col=0, text="C"),
                    TableCell(row=1, col=1, text="D"),
                ],
            ),
        ])
        doc = generator.generate_document(ir)
        table = doc.tables[0]
        assert len(table.rows) == 2
        assert len(table.columns) == 2
        assert table.cell(0, 0).text == "A"
        assert table.cell(1, 1).text == "D"

    def test_merged_cells(self, generator):
        ir = DocumentIR(body=[
            TableBlock(
                num_rows=2,
                num_cols=3,
                cells=[
                    TableCell(row=0, col=0, text="Span 2 cols", col_span=2),
                    TableCell(row=0, col=2, text="C"),
                    TableCell(row=1, col=0, text="D"),
                    TableCell(row=1, col=1, text="E"),
                    TableCell(row=1, col=2, text="F"),
                ],
            ),
        ])
        doc = generator.generate_document(ir)
        table = doc.tables[0]
        assert len(table.rows) == 2
        # Merged cell should contain the text
        assert table.cell(0, 0).text == "Span 2 cols"

    def test_empty_table(self, generator):
        ir = DocumentIR(body=[
            TableBlock(num_rows=0, num_cols=0, cells=[]),
        ])
        doc = generator.generate_document(ir)
        # Should still produce a table (1x1 placeholder)
        assert len(doc.tables) == 1


# ---------------------------------------------------------------------------
# Figures
# ---------------------------------------------------------------------------

class TestFigures:
    def test_missing_image_placeholder(self, generator):
        ir = DocumentIR(body=[
            FigureBlock(image_path="nonexistent.png"),
        ])
        doc = generator.generate_document(ir)
        # Should render a placeholder paragraph
        assert any("[Image not available]" in p.text for p in doc.paragraphs)

    def test_figure_with_caption(self, generator):
        ir = DocumentIR(body=[
            FigureBlock(image_path="nonexistent.png", caption="Figure 1"),
        ])
        doc = generator.generate_document(ir)
        texts = [p.text for p in doc.paragraphs]
        # Should have placeholder + caption
        assert "[Image not available]" in texts
        assert "Figure 1" in texts

    def test_image_from_file(self, generator, tmp_path):
        """Test embedding an actual image file."""
        # Create a small test PNG
        from PIL import Image

        img = Image.new("RGB", (100, 50), color="red")
        img_path = tmp_path / "test.png"
        img.save(str(img_path))

        ir = DocumentIR(body=[
            FigureBlock(image_path=str(img_path), caption="Test image"),
        ])
        doc = generator.generate_document(ir, base_dir=tmp_path)
        # Document should have inline shapes (the image)
        # and a caption paragraph
        assert any("Test image" in p.text for p in doc.paragraphs)

    def test_image_from_base64(self, generator):
        """Test embedding an image from base64."""
        import base64
        from PIL import Image
        import io

        img = Image.new("RGB", (100, 50), color="blue")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode()

        ir = DocumentIR(body=[
            FigureBlock(image_base64=b64),
        ])
        doc = generator.generate_document(ir)
        # Should successfully embed (no placeholder)
        texts = [p.text for p in doc.paragraphs]
        assert "[Image not available]" not in texts


# ---------------------------------------------------------------------------
# Page breaks
# ---------------------------------------------------------------------------

class TestPageBreaks:
    def test_page_break(self, generator):
        ir = DocumentIR(body=[
            ParagraphBlock(text="Before"),
            PageBreakBlock(),
            ParagraphBlock(text="After"),
        ])
        doc = generator.generate_document(ir)
        texts = [p.text for p in doc.paragraphs]
        # Should have: "Before", "" (page break para), "After"
        assert "Before" in texts
        assert "After" in texts


# ---------------------------------------------------------------------------
# Full document
# ---------------------------------------------------------------------------

class TestFullDocument:
    def test_complex_document(self, generator, tmp_path):
        """Generate a complex document and verify it saves correctly."""
        ir = DocumentIR(
            metadata=DocumentMetadata(
                source_file="test.pdf",
                title="Complex Doc",
                page_count=3,
            ),
            body=[
                HeadingBlock(level=1, text="Chapter 1", children=[
                    ParagraphBlock(
                        text="Intro with formatting",
                        runs=[
                            TextRun(text="Intro with "),
                            TextRun(text="formatting", bold=True, italic=True),
                        ],
                    ),
                    HeadingBlock(level=2, text="Section 1.1", children=[
                        ListBlock(style="ordered", items=[
                            ListItem(text="Step 1"),
                            ListItem(text="Step 2", children=[
                                ListItem(text="Sub-step 2a"),
                            ]),
                        ]),
                        TableBlock(
                            num_rows=2, num_cols=2,
                            cells=[
                                TableCell(row=0, col=0, text="ID"),
                                TableCell(row=0, col=1, text="Requirement"),
                                TableCell(row=1, col=0, text="REQ-001"),
                                TableCell(row=1, col=1, text="The system shall..."),
                            ],
                        ),
                    ]),
                ]),
                PageBreakBlock(),
                HeadingBlock(level=1, text="Chapter 2", children=[
                    ParagraphBlock(text="More content here."),
                ]),
            ],
        )

        out = tmp_path / "complex.docx"
        generator.generate(ir, out)
        assert out.exists()

        # Re-read and verify structure
        doc = open_docx(str(out))
        texts = [p.text for p in doc.paragraphs]

        assert "Chapter 1" in texts
        assert "Section 1.1" in texts
        assert "Step 1" in texts
        assert "Chapter 2" in texts
        assert "More content here." in texts

        # Verify table
        assert len(doc.tables) == 1
        assert doc.tables[0].cell(0, 0).text == "ID"
        assert doc.tables[0].cell(1, 1).text == "The system shall..."


class TestLowConfidenceMarking:
    def test_no_marking_by_default(self, tmp_path):
        """Low-confidence headings are NOT marked when mark_low_confidence=False."""
        ir = DocumentIR(body=[
            HeadingBlock(level=1, text="Iffy heading", confidence=0.40),
        ])
        generator = WordGenerator()
        out = tmp_path / "out.docx"
        generator.generate(ir, out)
        doc = open_docx(out)
        para = doc.paragraphs[0]
        # Should have no percentage marker
        assert not any("%" in r.text for r in para.runs)

    def test_marking_when_enabled(self, tmp_path):
        """Low-confidence headings get highlighted when enabled."""
        config = Config.default()
        config.style.mark_low_confidence = True
        config.style.low_confidence_threshold = 0.7
        ir = DocumentIR(body=[
            HeadingBlock(level=1, text="Iffy heading", confidence=0.40),
        ])
        generator = WordGenerator(config)
        out = tmp_path / "out.docx"
        generator.generate(ir, out)
        doc = open_docx(out)
        para = doc.paragraphs[0]
        # Should have highlight on first run
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        first_run = para.runs[0]
        rPr = first_run._element.rPr
        assert rPr is not None
        highlights = rPr.findall(f"{ns}highlight")
        assert len(highlights) == 1
        # Should have a percentage marker run
        assert any("[40%]" in r.text for r in para.runs)

    def test_high_confidence_not_marked(self, tmp_path):
        """Headings above the threshold are NOT marked even when enabled."""
        config = Config.default()
        config.style.mark_low_confidence = True
        ir = DocumentIR(body=[
            HeadingBlock(level=1, text="Good heading", confidence=0.95),
        ])
        generator = WordGenerator(config)
        out = tmp_path / "out.docx"
        generator.generate(ir, out)
        doc = open_docx(out)
        para = doc.paragraphs[0]
        assert not any("%" in r.text for r in para.runs)


class TestSaveErrors:
    def test_save_to_invalid_path(self, generator, simple_ir):
        with pytest.raises(Exception):
            generator.generate(simple_ir, Path("/nonexistent/dir/out.docx"))
